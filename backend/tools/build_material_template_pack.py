"""Build the reviewed material-template pack without expanding the raw archive.

The input is a traditional split ZIP (``.z01`` ... ``.zip``).  This module
parses the central directory and streams only allow-listed entries across disk
boundaries.  It intentionally does not concatenate or extract the archive.

The produced DOCX files are security-normalized OOXML packages.  They are
reference fixtures for a student demo and are not official government forms.
"""

from __future__ import annotations

import argparse
import binascii
import hashlib
import io
import json
import os
from pathlib import Path, PurePosixPath
import re
import struct
import tempfile
from typing import BinaryIO, Iterable, Iterator
import xml.etree.ElementTree as ET
import zipfile
import zlib


EOCD = b"PK\x05\x06"
CENTRAL = b"PK\x01\x02"
LOCAL = b"PK\x03\x04"
MAX_SOURCE_BYTES = 32 * 1024 * 1024
MAX_PACKAGE_BYTES = 64 * 1024 * 1024
MAX_PACKAGE_PARTS = 512
DEMO_NOTICE = "演示模板，仅供本项目填写演示，不作为正式政务表格。"

REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
DCTERMS_NS = "http://purl.org/dc/terms/"
EP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class PackBuildError(RuntimeError):
    """A reviewed source failed a deterministic safety or integrity gate."""


class _VirtualSplitFile:
    def __init__(self, parts: list[Path]):
        if not parts:
            raise PackBuildError("split ZIP has no parts")
        self.parts = parts
        self.sizes = [part.stat().st_size for part in parts]
        self.starts: list[int] = []
        cursor = 0
        for size in self.sizes:
            self.starts.append(cursor)
            cursor += size
        self.size = cursor

    def disk_offset(self, disk: int, offset: int) -> int:
        if disk < 0 or disk >= len(self.parts):
            raise PackBuildError(f"invalid split ZIP disk: {disk}")
        if offset < 0 or offset > self.sizes[disk]:
            raise PackBuildError(f"invalid offset {offset} on disk {disk}")
        return self.starts[disk] + offset

    def read_at(self, position: int, size: int) -> bytes:
        return b"".join(self.iter_at(position, size))

    def iter_at(self, position: int, size: int, chunk_size: int = 64 * 1024) -> Iterator[bytes]:
        if position < 0 or size < 0 or position + size > self.size:
            raise PackBuildError("split ZIP range exceeds archive")
        remaining = size
        while remaining:
            disk = max(index for index, start in enumerate(self.starts) if start <= position)
            offset = position - self.starts[disk]
            available = self.sizes[disk] - offset
            take = min(remaining, available, chunk_size)
            if take <= 0:
                raise PackBuildError("invalid zero-length split ZIP range")
            with self.parts[disk].open("rb") as stream:
                stream.seek(offset)
                block = stream.read(take)
            if len(block) != take:
                raise PackBuildError("truncated split ZIP part")
            yield block
            position += take
            remaining -= take


class SplitZipEntry:
    __slots__ = (
        "name",
        "raw_name",
        "flags",
        "method",
        "crc32",
        "compressed_size",
        "uncompressed_size",
        "disk_start",
        "local_offset",
    )

    def __init__(
        self,
        *,
        name: str,
        raw_name: bytes,
        flags: int,
        method: int,
        crc32: int,
        compressed_size: int,
        uncompressed_size: int,
        disk_start: int,
        local_offset: int,
    ) -> None:
        self.name = name
        self.raw_name = raw_name
        self.flags = flags
        self.method = method
        self.crc32 = crc32
        self.compressed_size = compressed_size
        self.uncompressed_size = uncompressed_size
        self.disk_start = disk_start
        self.local_offset = local_offset


class SplitZipArchive:
    """Minimal read-only reader for non-ZIP64 traditional split archives."""

    def __init__(self, final_zip: Path):
        self.final_zip = final_zip.resolve(strict=True)
        stem = self.final_zip.with_suffix("")
        numbered: list[tuple[int, Path]] = []
        for path in self.final_zip.parent.glob(f"{stem.name}.z[0-9][0-9]"):
            match = re.fullmatch(r"\.z(\d\d)", path.suffix, flags=re.IGNORECASE)
            if match:
                numbered.append((int(match.group(1)), path.resolve(strict=True)))
        numbered.sort(key=lambda item: item[0])
        if not numbered or numbered[0][0] != 1:
            raise PackBuildError("expected traditional split ZIP starting at .z01")
        expected = list(range(1, numbered[-1][0] + 1))
        if [number for number, _ in numbered] != expected:
            raise PackBuildError("split ZIP part sequence has a gap")
        self.parts = [path for _, path in numbered] + [self.final_zip]
        self.virtual = _VirtualSplitFile(self.parts)
        self.entries = self._read_central_directory()

    def _read_central_directory(self) -> dict[str, SplitZipEntry]:
        last = self.parts[-1].read_bytes()
        offset = last.rfind(EOCD, max(0, len(last) - 65_557))
        if offset < 0 or offset + 22 > len(last):
            raise PackBuildError("end-of-central-directory record not found")
        (
            signature,
            disk_number,
            central_disk,
            disk_entries,
            total_entries,
            central_size,
            central_offset,
            comment_length,
        ) = struct.unpack_from("<4s4H2LH", last, offset)
        if signature != EOCD or offset + 22 + comment_length != len(last):
            raise PackBuildError("invalid end-of-central-directory record")
        if disk_number != len(self.parts) - 1 or central_disk >= len(self.parts):
            raise PackBuildError("split ZIP disk metadata does not match available parts")
        if disk_entries != total_entries:
            raise PackBuildError("central directory spans disks; unsupported archive layout")
        if central_size == 0xFFFFFFFF or central_offset == 0xFFFFFFFF:
            raise PackBuildError("ZIP64 is not supported")
        position = self.virtual.disk_offset(central_disk, central_offset)
        end = position + central_size
        entries: dict[str, SplitZipEntry] = {}
        for _ in range(total_entries):
            fixed = self.virtual.read_at(position, 46)
            values = struct.unpack("<4s6H3L5H2L", fixed)
            if values[0] != CENTRAL:
                raise PackBuildError("invalid central-directory signature")
            (
                _,
                _made_by,
                _needed,
                flags,
                method,
                _mtime,
                _mdate,
                crc32,
                compressed_size,
                uncompressed_size,
                name_length,
                extra_length,
                comment_length,
                disk_start,
                _internal,
                _external,
                local_offset,
            ) = values
            if 0xFFFFFFFF in (compressed_size, uncompressed_size, local_offset) or disk_start == 0xFFFF:
                raise PackBuildError("ZIP64 entry is not supported")
            variable = self.virtual.read_at(position + 46, name_length + extra_length + comment_length)
            raw_name = variable[:name_length]
            extra = variable[name_length : name_length + extra_length]
            name = _decode_zip_name(raw_name, flags, extra)
            _validate_archive_name(name)
            if name in entries:
                raise PackBuildError(f"duplicate archive entry: {name}")
            entries[name] = SplitZipEntry(
                name=name,
                raw_name=raw_name,
                flags=flags,
                method=method,
                crc32=crc32,
                compressed_size=compressed_size,
                uncompressed_size=uncompressed_size,
                disk_start=disk_start,
                local_offset=local_offset,
            )
            position += 46 + name_length + extra_length + comment_length
        if position != end:
            raise PackBuildError("central-directory size mismatch")
        return entries

    def iter_entry(self, name: str, *, max_size: int = MAX_SOURCE_BYTES) -> Iterator[bytes]:
        try:
            entry = self.entries[name]
        except KeyError as exc:
            raise PackBuildError(f"allow-listed source not found: {name}") from exc
        if entry.flags & 0x1:
            raise PackBuildError(f"encrypted source is forbidden: {name}")
        if entry.method not in (0, 8):
            raise PackBuildError(f"unsupported compression method {entry.method}: {name}")
        if entry.uncompressed_size > max_size:
            raise PackBuildError(f"source exceeds {max_size} bytes: {name}")
        local_position = self.virtual.disk_offset(entry.disk_start, entry.local_offset)
        fixed = self.virtual.read_at(local_position, 30)
        values = struct.unpack("<4s5H3L2H", fixed)
        if values[0] != LOCAL:
            raise PackBuildError(f"invalid local header: {name}")
        local_flags = values[2]
        local_method = values[3]
        name_length, extra_length = values[-2:]
        if local_flags != entry.flags or local_method != entry.method:
            raise PackBuildError(f"central/local metadata mismatch: {name}")
        raw_local_name = self.virtual.read_at(local_position + 30, name_length)
        if raw_local_name != entry.raw_name:
            raise PackBuildError(f"central/local filename mismatch: {name}")
        data_position = local_position + 30 + name_length + extra_length
        inflater = zlib.decompressobj(-15) if entry.method == 8 else None
        digest_crc = 0
        produced = 0
        for compressed in self.virtual.iter_at(data_position, entry.compressed_size):
            block = inflater.decompress(compressed) if inflater else compressed
            if block:
                produced += len(block)
                if produced > max_size:
                    raise PackBuildError(f"expanded source exceeds limit: {name}")
                digest_crc = binascii.crc32(block, digest_crc)
                yield block
        if inflater:
            block = inflater.flush()
            if block:
                produced += len(block)
                digest_crc = binascii.crc32(block, digest_crc)
                yield block
            if not inflater.eof:
                raise PackBuildError(f"truncated deflate stream: {name}")
        if produced != entry.uncompressed_size or digest_crc & 0xFFFFFFFF != entry.crc32:
            raise PackBuildError(f"size or CRC mismatch: {name}")

    def read_entry(self, name: str, *, max_size: int = MAX_SOURCE_BYTES) -> bytes:
        return b"".join(self.iter_entry(name, max_size=max_size))


def _decode_zip_name(raw_name: bytes, flags: int, extra: bytes) -> str:
    cursor = 0
    while cursor + 4 <= len(extra):
        field_id, size = struct.unpack_from("<HH", extra, cursor)
        cursor += 4
        value = extra[cursor : cursor + size]
        cursor += size
        if field_id == 0x7075 and len(value) >= 5 and value[0] == 1:
            if struct.unpack_from("<L", value, 1)[0] == binascii.crc32(raw_name) & 0xFFFFFFFF:
                return value[5:].decode("utf-8")
    if flags & 0x800:
        return raw_name.decode("utf-8")
    # Some Windows ZIP writers store UTF-8 names without setting bit 11.  Try
    # strict UTF-8 first, then the common Simplified Chinese legacy encoding.
    for encoding in ("utf-8", "gb18030", "cp437"):
        try:
            return raw_name.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise PackBuildError("archive filename cannot be decoded")


def _validate_archive_name(name: str) -> None:
    normalized = name.replace("\\", "/")
    pure = PurePosixPath(normalized)
    if (
        not name
        or "\x00" in name
        or pure.is_absolute()
        or any(part in ("", ".", "..") for part in pure.parts)
        or re.match(r"^[A-Za-z]:", normalized)
    ):
        raise PackBuildError(f"unsafe archive path: {name!r}")


def _safe_docx_parts(data: bytes, source_name: str) -> dict[str, bytes]:
    if len(data) > MAX_SOURCE_BYTES:
        raise PackBuildError(f"DOCX source is too large: {source_name}")
    try:
        archive = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile as exc:
        raise PackBuildError(f"source is not a valid DOCX: {source_name}") from exc
    infos = archive.infolist()
    if len(infos) > MAX_PACKAGE_PARTS:
        raise PackBuildError(f"DOCX contains too many package parts: {source_name}")
    total = 0
    parts: dict[str, bytes] = {}
    forbidden_markers = (
        "vbaproject",
        "activex/",
        "embeddings/",
        "altchunk",
        "oleobject",
    )
    for info in infos:
        name = info.filename.replace("\\", "/")
        _validate_archive_name(name)
        lowered = name.lower()
        if info.flag_bits & 0x1:
            raise PackBuildError(f"encrypted DOCX part is forbidden: {name}")
        if any(marker in lowered for marker in forbidden_markers) or lowered.endswith(".bin"):
            raise PackBuildError(f"active or opaque DOCX content is forbidden: {name}")
        total += info.file_size
        if total > MAX_PACKAGE_BYTES:
            raise PackBuildError(f"DOCX expanded package exceeds limit: {source_name}")
        if name in parts:
            raise PackBuildError(f"duplicate DOCX package part: {name}")
        parts[name] = archive.read(info)
    for required in ("[Content_Types].xml", "_rels/.rels", "word/document.xml"):
        if required not in parts:
            raise PackBuildError(f"DOCX lacks required part {required}: {source_name}")
    for name, payload in parts.items():
        if name.endswith(".rels"):
            root = ET.fromstring(payload)
            for relationship in root:
                if relationship.get("TargetMode", "").lower() == "external":
                    raise PackBuildError(f"external relationship is forbidden in {source_name}: {name}")
    content_types = parts["[Content_Types].xml"].lower()
    if b"macroenabled" in content_types or b"vba" in content_types:
        raise PackBuildError(f"macro-enabled DOCX is forbidden: {source_name}")
    return parts


def sanitize_docx(data: bytes, source_name: str) -> bytes:
    parts = _safe_docx_parts(data, source_name)
    for part_name in tuple(parts):
        if part_name.lower().startswith("customxml/"):
            parts.pop(part_name)
    parts.pop("docProps/custom.xml", None)
    parts.pop("docProps/thumbnail.jpeg", None)

    for rel_name, rel_payload in list(parts.items()):
        if not rel_name.endswith(".rels"):
            continue
        relationships = ET.fromstring(rel_payload)
        changed = False
        for child in list(relationships):
            target = child.get("Target", "").replace("\\", "/").lower()
            rel_type = child.get("Type", "").lower()
            if (
                target in ("docprops/custom.xml", "docprops/thumbnail.jpeg")
                or "customxml/" in target
                or rel_type.endswith("/custom-properties")
                or rel_type.endswith("/customxml")
            ):
                relationships.remove(child)
                changed = True
        if changed:
            parts[rel_name] = ET.tostring(relationships, encoding="utf-8", xml_declaration=True)

    content_root = ET.fromstring(parts["[Content_Types].xml"])
    for child in list(content_root):
        part_name = child.get("PartName", "").lstrip("/").lower()
        if part_name in ("docprops/custom.xml", "docprops/thumbnail.jpeg") or part_name.startswith(
            "customxml/"
        ):
            content_root.remove(child)
    parts["[Content_Types].xml"] = ET.tostring(content_root, encoding="utf-8", xml_declaration=True)

    if "docProps/core.xml" in parts:
        core = ET.fromstring(parts["docProps/core.xml"])
        for tag in (
            f"{{{DC_NS}}}creator",
            f"{{{CP_NS}}}lastModifiedBy",
            f"{{{DC_NS}}}subject",
            f"{{{DC_NS}}}description",
            f"{{{CP_NS}}}keywords",
            f"{{{CP_NS}}}category",
        ):
            node = core.find(tag)
            if node is not None:
                node.text = ""
        for tag in (f"{{{DCTERMS_NS}}}created", f"{{{DCTERMS_NS}}}modified"):
            node = core.find(tag)
            if node is not None:
                core.remove(node)
        parts["docProps/core.xml"] = ET.tostring(core, encoding="utf-8", xml_declaration=True)

    if "docProps/app.xml" in parts:
        app = ET.fromstring(parts["docProps/app.xml"])
        for local in ("Company", "Manager", "HyperlinkBase"):
            node = app.find(f"{{{EP_NS}}}{local}")
            if node is not None:
                node.text = ""
        parts["docProps/app.xml"] = ET.tostring(app, encoding="utf-8", xml_declaration=True)

    rsid_pattern = re.compile(rb"\s+w:rsid(?:R|RPr|RDefault|P|Del|Sect|Tr|RPr|Root)?=\"[^\"]*\"")
    for name, payload in list(parts.items()):
        if name.startswith("word/") and name.endswith(".xml"):
            parts[name] = rsid_pattern.sub(b"", payload)

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for name in sorted(parts):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            archive.writestr(info, parts[name])
    normalized = output.getvalue()
    _safe_docx_parts(normalized, source_name)
    return normalized


def _set_cell_text(cell, text: str, *, bold: bool = False, centered: bool = False) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def _set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(sum(widths_dxa)))
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(indent_dxa))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa, strict=True):
            cell.width = value
            tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(value))


def _add_demo_footer(document) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(DEMO_NOTICE)
        run.font.name = "微软雅黑"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_demo_footer(data: bytes) -> bytes:
    from docx import Document

    document = Document(io.BytesIO(data))
    _add_demo_footer(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    output = io.BytesIO()
    document.save(output)
    return sanitize_docx(output.getvalue(), "footer-normalized DOCX")


def build_personnel_summary_fixture() -> bytes:
    """Create the deterministic VISUAL_RECONSTRUCT table fixture."""
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run("缴存人员汇总表（演示模板）")
    title_run.bold = True
    title_run.font.name = "微软雅黑"
    title_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)

    meta = document.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    meta_values = (
        ("单位名称", "【待填写】", "统一社会信用代码", "【待填写】"),
        ("经办人", "【待填写】", "联系电话", "【待填写】"),
    )
    for row, values in zip(meta.rows, meta_values, strict=True):
        for index, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
            _set_cell_text(cell, value, bold=index % 2 == 0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_table_geometry(meta, [1400, 2880, 1800, 2880])

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)

    table = document.add_table(rows=7, cols=7)
    table.style = "Table Grid"
    headers = ("序号", "姓名", "证件号码（演示）", "缴存基数", "单位比例", "个人比例", "备注")
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_text(cell, header, bold=True, centered=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows[1:], start=1):
        values = (str(row_index), "【待填写】", "【待填写】", "【待填写】", "【待填写】", "【待填写】", "")
        for column_index, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
            _set_cell_text(cell, value, centered=column_index != 1)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_table_geometry(table, [600, 1100, 2260, 1300, 1200, 1200, 1500])

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run("填写说明：仅填写合成演示数据；人员行可按需复制增加。")
    note_run.font.name = "宋体"
    note_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _add_demo_footer(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    output = io.BytesIO()
    document.save(output)
    return sanitize_docx(output.getvalue(), "generated personnel-summary fixture")


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as stream:
        value = json.load(stream)
    if not isinstance(value, dict):
        raise PackBuildError(f"expected JSON object: {path}")
    return value


def build_pack(final_zip: Path, source_map_path: Path, output_dir: Path) -> dict:
    archive = SplitZipArchive(final_zip)
    source_map = _load_json(source_map_path)
    templates = source_map.get("templates")
    if source_map.get("version") != 1 or not isinstance(templates, list):
        raise PackBuildError("source map must contain version=1 and templates[]")
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_dir = output_dir / "docx"
    docx_dir.mkdir(parents=True, exist_ok=True)
    source_cache: dict[str, tuple[bytes, str]] = {}
    packaged_cache: dict[str, tuple[bytes, str]] = {}
    built_templates: list[dict] = []

    for item in templates:
        if not isinstance(item, dict):
            raise PackBuildError("template source-map entry must be an object")
        result = dict(item)
        source_entry = item.get("source_archive_entry")
        packaged_path = item.get("packaged_path")
        mode = item.get("mode")
        if mode not in ("SOURCE_EDITABLE", "VISUAL_RECONSTRUCT", "NOT_GENERATABLE"):
            raise PackBuildError(f"invalid template mode: {mode}")
        if source_entry is not None:
            if not isinstance(source_entry, str):
                raise PackBuildError("source_archive_entry must be a string or null")
            if source_entry not in source_cache:
                raw = archive.read_entry(source_entry)
                source_cache[source_entry] = (raw, sha256_bytes(raw))
            source_bytes, source_hash = source_cache[source_entry]
            result["source_sha256"] = source_hash
        else:
            source_bytes = None
            result["source_sha256"] = None

        if packaged_path is not None:
            if not isinstance(packaged_path, str) or PurePosixPath(packaged_path).is_absolute():
                raise PackBuildError("packaged_path must be a safe relative path")
            _validate_archive_name(packaged_path)
            destination = (output_dir / Path(packaged_path)).resolve()
            if output_dir.resolve() not in destination.parents:
                raise PackBuildError("packaged_path escapes output directory")
            cache_key = packaged_path
            if cache_key not in packaged_cache:
                fixture_kind = item.get("fixture_kind")
                if fixture_kind == "PERSONNEL_SUMMARY":
                    packaged = build_personnel_summary_fixture()
                elif source_bytes is not None:
                    packaged = add_demo_footer(sanitize_docx(source_bytes, source_entry))
                else:
                    raise PackBuildError(f"packaged template has no source or fixture: {packaged_path}")
                packaged_cache[cache_key] = (packaged, sha256_bytes(packaged))
                destination.parent.mkdir(parents=True, exist_ok=True)
                destination.write_bytes(packaged)
            _, packaged_hash = packaged_cache[cache_key]
            result["packaged_sha256"] = packaged_hash
        else:
            result["packaged_sha256"] = None
        result.pop("fixture_kind", None)
        built_templates.append(result)

    manifest = {"version": 1, "templates": built_templates}
    manifest_path = output_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest


def list_entries(final_zip: Path, pattern: str | None) -> None:
    archive = SplitZipArchive(final_zip)
    matcher = re.compile(pattern, flags=re.IGNORECASE) if pattern else None
    for name in sorted(archive.entries):
        if matcher is None or matcher.search(name):
            entry = archive.entries[name]
            print(f"{entry.uncompressed_size:>10}\t{name}")


def verify_pack(manifest_path: Path) -> None:
    manifest = _load_json(manifest_path)
    templates = manifest.get("templates")
    if manifest.get("version") != 1 or not isinstance(templates, list):
        raise PackBuildError("invalid generated manifest")
    seen_ids: set[str] = set()
    seen_requirements: set[tuple[str, str]] = set()
    root = manifest_path.parent.resolve()
    for item in templates:
        template_id = item.get("template_id")
        if not isinstance(template_id, str) or not re.fullmatch(r"[a-z0-9][a-z0-9-]{2,63}", template_id):
            raise PackBuildError(f"invalid template_id: {template_id}")
        if template_id in seen_ids:
            raise PackBuildError(f"duplicate template_id: {template_id}")
        seen_ids.add(template_id)
        key = (item.get("service_code"), item.get("requirement_code"))
        if not all(isinstance(part, str) and part for part in key) or key in seen_requirements:
            raise PackBuildError(f"duplicate/invalid service requirement: {key}")
        seen_requirements.add(key)
        if not isinstance(item.get("allowed_fields"), list) or not all(
            isinstance(field, str) and re.fullmatch(r"[a-z][a-z0-9_]*(?:\.[a-z][a-z0-9_]*)*", field)
            for field in item["allowed_fields"]
        ):
            raise PackBuildError(f"invalid allowed_fields: {template_id}")
        path = item.get("packaged_path")
        if path is None:
            if item.get("packaged_sha256") is not None:
                raise PackBuildError(f"hash without packaged_path: {template_id}")
            continue
        destination = (root / path).resolve()
        if root not in destination.parents or not destination.is_file():
            raise PackBuildError(f"missing packaged DOCX: {template_id}")
        data = destination.read_bytes()
        if sha256_bytes(data) != item.get("packaged_sha256"):
            raise PackBuildError(f"packaged hash mismatch: {template_id}")
        sanitize_docx(data, f"pack verification {template_id}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)
    list_parser = subparsers.add_parser("list", help="list central-directory entries without extracting")
    list_parser.add_argument("--archive", type=Path, required=True, help="path to final .zip part")
    list_parser.add_argument("--pattern", help="optional regular expression")
    build_parser = subparsers.add_parser("build", help="build the reviewed template pack")
    build_parser.add_argument("--archive", type=Path, required=True, help="path to final .zip part")
    build_parser.add_argument("--source-map", type=Path, required=True)
    build_parser.add_argument("--output-dir", type=Path, required=True)
    verify_parser = subparsers.add_parser("verify", help="verify a built manifest and every packaged DOCX")
    verify_parser.add_argument("--manifest", type=Path, required=True)
    args = parser.parse_args()
    if args.command == "list":
        list_entries(args.archive, args.pattern)
    elif args.command == "build":
        build_pack(args.archive, args.source_map, args.output_dir)
    elif args.command == "verify":
        verify_pack(args.manifest)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
