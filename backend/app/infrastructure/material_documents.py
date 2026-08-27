from __future__ import annotations

import asyncio
import base64
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path, PurePosixPath
from typing import Any
from xml.etree import ElementTree

import httpx
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.application.material_documents import (
    MaterialDocumentProcessingError,
    TemplateAnalysis,
)
from app.domain.enums import MaterialTemplateMode


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_NOTICE = "演示模板，仅供项目填写演示，不作为正式政务表格。"


_SAFE_MATERIAL_TEMPLATE_ERROR_CODES = frozenset(
    {
        "docx_active_content",
        "docx_altchunk",
        "docx_entries",
        "docx_expansion",
        "docx_external_relationship",
        "docx_invalid",
        "docx_path",
        "docx_size",
        "docx_structure",
        "material_template_error",
        "material_documents_disabled",
        "model_json_invalid",
        "model_key_missing",
        "model_request_failed",
        "model_schema_field",
        "model_schema_fields",
        "model_schema_invalid",
        "model_schema_value",
        "model_schema_warning",
        "model_schema_warnings",
        "template_layout_unsupported",
        "template_manifest_duplicate",
        "template_manifest_field",
        "template_manifest_fields",
        "template_manifest_invalid",
        "template_manifest_item",
        "template_manifest_items",
        "template_manifest_missing",
        "template_manifest_mode",
        "template_manifest_path",
        "template_manifest_sha256",
        "template_manifest_source",
        "template_manifest_source_missing",
        "template_manifest_source_unreadable",
        "template_manifest_text",
        "template_manifest_token",
        "template_manifest_version",
        "template_not_generatable",
        "template_pack_sources_empty",
        "template_page_count",
        "template_page_output",
        "template_page_size",
        "template_pages_empty",
        "template_pages_size",
        "template_pdf_empty",
        "template_renderer_unavailable",
    }
)


class MaterialTemplateError(MaterialDocumentProcessingError):
    """Template failure with an allowlisted code safe for durable job state."""

    def __init__(self, code: str) -> None:
        super().__init__(
            code
            if code in _SAFE_MATERIAL_TEMPLATE_ERROR_CODES
            else "material_template_error"
        )


@dataclass(frozen=True, slots=True)
class PackedMaterialTemplate:
    template_key: str
    service_code: str
    requirement_code: str
    title: str
    mode: MaterialTemplateMode
    source_object_key: str | None
    source_sha256: str | None
    source_bytes: bytes | None
    allowed_fields: tuple[str, ...]
    notice: str


class MaterialTemplatePack:
    """Loads only the reviewed, repository-baked subset of the source archive."""

    def __init__(self, manifest_path: str, source_prefix: str) -> None:
        path = Path(manifest_path)
        if not path.is_absolute():
            backend_root = Path(__file__).resolve().parents[2]
            path = backend_root / path
        self.manifest_path = path.resolve()
        self.source_prefix = source_prefix.strip("/")

    def load(self) -> tuple[PackedMaterialTemplate, ...]:
        try:
            payload = json.loads(self.manifest_path.read_text(encoding="utf-8"))
        except FileNotFoundError as exc:
            raise MaterialTemplateError("template_manifest_missing") from exc
        except Exception as exc:
            raise MaterialTemplateError("template_manifest_invalid") from exc
        if not isinstance(payload, dict) or payload.get("version") != 1:
            raise MaterialTemplateError("template_manifest_version")
        raw_templates = payload.get("templates")
        if (
            not isinstance(raw_templates, list)
            or not raw_templates
            or len(raw_templates) > 64
        ):
            raise MaterialTemplateError("template_manifest_items")
        seen: set[str] = set()
        result: list[PackedMaterialTemplate] = []
        root = self.manifest_path.parent
        for raw in raw_templates:
            if not isinstance(raw, dict):
                raise MaterialTemplateError("template_manifest_item")
            template_key = self._token(raw.get("template_id"), 96)
            service_code = self._token(raw.get("service_code"), 64)
            requirement_code = self._token(raw.get("requirement_code"), 64)
            if template_key in seen:
                raise MaterialTemplateError("template_manifest_duplicate")
            seen.add(template_key)
            try:
                mode = MaterialTemplateMode(str(raw.get("mode")))
            except ValueError as exc:
                raise MaterialTemplateError("template_manifest_mode") from exc
            title = self._text(raw.get("title"), 200)
            notice = self._text(raw.get("notice") or _NOTICE, 500)
            allowed = raw.get("allowed_fields", [])
            if (
                not isinstance(allowed, list)
                or len(allowed) > 64
                or any(not isinstance(item, str) for item in allowed)
            ):
                raise MaterialTemplateError("template_manifest_fields")
            fields = tuple(self._field(item) for item in allowed)
            packaged = raw.get("packaged_path")
            source_bytes: bytes | None = None
            source_sha: str | None = None
            source_key: str | None = None
            if packaged:
                if not isinstance(packaged, str):
                    raise MaterialTemplateError("template_manifest_path")
                candidate = (root / packaged).resolve()
                try:
                    candidate.relative_to(root.resolve())
                except ValueError as exc:
                    raise MaterialTemplateError("template_manifest_path") from exc
                try:
                    source_bytes = candidate.read_bytes()
                except FileNotFoundError as exc:
                    raise MaterialTemplateError(
                        "template_manifest_source_missing"
                    ) from exc
                except OSError as exc:
                    raise MaterialTemplateError(
                        "template_manifest_source_unreadable"
                    ) from exc
                validate_docx_package(source_bytes)
                source_sha = hashlib.sha256(source_bytes).hexdigest()
                declared_sha = raw.get("packaged_sha256")
                if (
                    not isinstance(declared_sha, str)
                    or not re.fullmatch(r"[0-9a-f]{64}", declared_sha)
                    or declared_sha != source_sha
                ):
                    raise MaterialTemplateError("template_manifest_sha256")
                source_key = f"{self.source_prefix}/{template_key}.docx"
            if mode is not MaterialTemplateMode.NOT_GENERATABLE and source_bytes is None:
                raise MaterialTemplateError("template_manifest_source")
            result.append(
                PackedMaterialTemplate(
                    template_key,
                    service_code,
                    requirement_code,
                    title,
                    mode,
                    source_key,
                    source_sha,
                    source_bytes,
                    fields,
                    notice,
                )
            )
        return tuple(result)

    @staticmethod
    def _token(value: Any, maximum: int) -> str:
        if not isinstance(value, str) or not re.fullmatch(r"[A-Za-z0-9_.-]+", value):
            raise MaterialTemplateError("template_manifest_token")
        return value[:maximum]

    @staticmethod
    def _field(value: str) -> str:
        if not re.fullmatch(r"[A-Za-z0-9_.-]{1,128}", value):
            raise MaterialTemplateError("template_manifest_field")
        return value

    @staticmethod
    def _text(value: Any, maximum: int) -> str:
        if not isinstance(value, str) or not value.strip() or len(value) > maximum:
            raise MaterialTemplateError("template_manifest_text")
        return re.sub(r"\s+", " ", value).strip()


def validate_docx_package(content: bytes) -> None:
    if not content or len(content) > 20 * 1024 * 1024:
        raise MaterialTemplateError("docx_size")
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            infos = archive.infolist()
            if len(infos) > 2000:
                raise MaterialTemplateError("docx_entries")
            total = 0
            names: set[str] = set()
            for info in infos:
                path = PurePosixPath(info.filename)
                if path.is_absolute() or ".." in path.parts or "\\" in info.filename:
                    raise MaterialTemplateError("docx_path")
                total += info.file_size
                if total > 100 * 1024 * 1024:
                    raise MaterialTemplateError("docx_expansion")
                lower = info.filename.lower()
                if (
                    "vbaproject" in lower
                    or lower.startswith("word/embeddings/")
                    or lower.endswith(".bin")
                ):
                    raise MaterialTemplateError("docx_active_content")
                names.add(info.filename)
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise MaterialTemplateError("docx_structure")
            for name in names:
                if not name.endswith(".rels"):
                    continue
                root = ElementTree.fromstring(archive.read(name))
                for relationship in root:
                    if relationship.attrib.get("TargetMode") == "External":
                        raise MaterialTemplateError("docx_external_relationship")
            document_xml = archive.read("word/document.xml")
            if b"altChunk" in document_xml:
                raise MaterialTemplateError("docx_altchunk")
    except MaterialTemplateError:
        raise
    except Exception as exc:
        raise MaterialTemplateError("docx_invalid") from exc


class MockMaterialTemplateAnalyzer:
    model_name = "mock-material-template-v1"

    async def analyze(
        self,
        *,
        page_images: tuple[bytes, ...],
        template_title: str,
        allowed_fields: tuple[str, ...],
        form_snapshot: dict[str, Any],
        request_text: str | None,
    ) -> TemplateAnalysis:
        if not page_images:
            raise MaterialTemplateError("template_pages_empty")
        fields: dict[str, str] = {}
        for field in allowed_fields:
            value = _lookup(form_snapshot, field)
            if isinstance(value, (str, int, float, bool)) and str(value).strip():
                fields[field] = str(value).strip()[:500]
            elif isinstance(value, (list, dict)):
                encoded = json.dumps(value, ensure_ascii=False, separators=(",", ":"))
                if len(encoded) <= 4000:
                    fields[field] = encoded
        return TemplateAnalysis(fields=fields)

    async def close(self) -> None:
        return None


class DashScopeMaterialTemplateAnalyzer:
    base_url = "https://dashscope.aliyuncs.com/compatible-mode/v1"

    def __init__(
        self,
        api_key: str | None,
        *,
        model_name: str,
        base_url: str,
        timeout_seconds: float,
        client: httpx.AsyncClient | None = None,
    ) -> None:
        self._api_key = (api_key or "").strip()
        if base_url.strip().rstrip("/") != self.base_url:
            raise ValueError("unsupported DashScope material-template base URL")
        self.model_name = model_name
        self.endpoint = f"{self.base_url}/chat/completions"
        self._client = client or httpx.AsyncClient(
            timeout=httpx.Timeout(timeout_seconds), trust_env=False
        )
        self._owns_client = client is None

    async def analyze(
        self,
        *,
        page_images: tuple[bytes, ...],
        template_title: str,
        allowed_fields: tuple[str, ...],
        form_snapshot: dict[str, Any],
        request_text: str | None,
    ) -> TemplateAnalysis:
        if not self._api_key:
            raise MaterialTemplateError("model_key_missing")
        if not 1 <= len(page_images) <= 12:
            raise MaterialTemplateError("template_page_count")
        image_parts: list[dict[str, Any]] = []
        total = 0
        for image in page_images:
            if not image or len(image) > 2 * 1024 * 1024:
                raise MaterialTemplateError("template_page_size")
            total += len(image)
            image_parts.append(
                {
                    "type": "image_url",
                    "image_url": {
                        "url": "data:image/jpeg;base64,"
                        + base64.b64encode(image).decode("ascii")
                    },
                }
            )
        if total > 12 * 1024 * 1024:
            raise MaterialTemplateError("template_pages_size")
        snapshot = _bounded_snapshot(form_snapshot, allowed_fields)
        instruction = {
            "template_title": template_title,
            "allowed_fields": list(allowed_fields),
            "application_values": snapshot,
            "user_request": request_text or "",
            "output_contract": {
                "fields": "object; keys must be from allowed_fields; scalar string values only",
                "warnings": "optional array of at most 8 short strings",
            },
        }
        try:
            response = await self._client.post(
                self.endpoint,
                headers={"Authorization": f"Bearer {self._api_key}"},
                json={
                    "model": self.model_name,
                    "messages": [
                        {
                            "role": "system",
                            "content": (
                                "你是政务演示材料模板字段提取器。只分析版式与空白字段，"
                                "只使用给定的合成演示值；不得生成签名、盖章、证件号或身份事实。"
                                "必须只返回一个 JSON 对象，不能返回 OOXML、Markdown 或额外文字。"
                            ),
                        },
                        {
                            "role": "user",
                            "content": image_parts
                            + [
                                {
                                    "type": "text",
                                    "text": json.dumps(
                                        instruction, ensure_ascii=False
                                    ),
                                }
                            ],
                        },
                    ],
                    "response_format": {"type": "json_object"},
                    "enable_thinking": False,
                    "temperature": 0,
                },
                follow_redirects=False,
            )
            response.raise_for_status()
        except httpx.HTTPError as exc:
            raise MaterialTemplateError("model_request_failed") from exc
        try:
            content = response.json()["choices"][0]["message"]["content"]
            raw = json.loads(content)
        except Exception as exc:
            raise MaterialTemplateError("model_json_invalid") from exc
        return _decode_analysis(raw, allowed_fields)

    async def close(self) -> None:
        if self._owns_client:
            await self._client.aclose()


class LibreOfficeTemplatePageRenderer:
    def __init__(self, *, max_pages: int = 12) -> None:
        self.max_pages = max_pages

    async def render_pages(self, source_docx: bytes) -> tuple[bytes, ...]:
        validate_docx_package(source_docx)
        return await asyncio.to_thread(self._render, source_docx)

    def _render(self, source_docx: bytes) -> tuple[bytes, ...]:
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        pdftoppm = shutil.which("pdftoppm")
        if not soffice or not pdftoppm:
            raise MaterialTemplateError("template_renderer_unavailable")
        with tempfile.TemporaryDirectory(prefix="material-template-") as directory:
            root = Path(directory)
            source = root / "source.docx"
            # The reviewed pack is deliberately rewritten by python-docx to
            # remove unsupported producer quirks before LibreOffice opens it.
            # Some otherwise valid, minimal OOXML packages are accepted by
            # Word/python-docx but LibreOffice exits 0 without producing a PDF.
            # Normalizing the already hash-verified package makes that failure
            # deterministic without changing the immutable source object.
            Document(io.BytesIO(source_docx)).save(source)
            profile = root / "lo-profile"
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    f"-env:UserInstallation={profile.as_uri()}",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(root),
                    str(source),
                ],
                check=True,
                timeout=60,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                env={**os.environ, "HOME": str(root)},
            )
            pdf = root / "source.pdf"
            if not pdf.is_file() or pdf.stat().st_size == 0:
                raise MaterialTemplateError("template_pdf_empty")
            prefix = root / "page"
            subprocess.run(
                [
                    pdftoppm,
                    "-f",
                    "1",
                    "-l",
                    str(self.max_pages),
                    "-r",
                    "144",
                    "-jpeg",
                    "-jpegopt",
                    "quality=85",
                    str(pdf),
                    str(prefix),
                ],
                check=True,
                timeout=60,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            pages = tuple(path.read_bytes() for path in sorted(root.glob("page-*.jpg")))
            if not pages or any(len(page) > 2 * 1024 * 1024 for page in pages):
                raise MaterialTemplateError("template_page_output")
            return pages


class DeterministicDocxMaterialRenderer:
    async def render(
        self,
        *,
        mode: MaterialTemplateMode,
        template_key: str,
        template_title: str,
        source_docx: bytes,
        analysis: TemplateAnalysis,
    ) -> bytes:
        return await asyncio.to_thread(
            self._render, mode, template_key, template_title, source_docx, analysis
        )

    def _render(
        self,
        mode: MaterialTemplateMode,
        template_key: str,
        template_title: str,
        source_docx: bytes,
        analysis: TemplateAnalysis,
    ) -> bytes:
        validate_docx_package(source_docx)
        if mode is MaterialTemplateMode.SOURCE_EDITABLE:
            document = Document(io.BytesIO(source_docx))
            matched = self._replace_placeholders(document, analysis.fields)
            self._fill_labeled_blanks(document, analysis.fields, matched)
        elif mode is MaterialTemplateMode.VISUAL_RECONSTRUCT:
            document = self._reconstruct(
                template_key, template_title, analysis.fields
            )
        else:
            raise MaterialTemplateError("template_not_generatable")
        self._add_notice_footer(document)
        self._scrub_properties(document)
        buffer = io.BytesIO()
        document.save(buffer)
        result = buffer.getvalue()
        validate_docx_package(result)
        return result

    def _reconstruct(
        self, template_key: str, title_text: str, fields: dict[str, str]
    ) -> Any:
        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(title_text)
        run.bold = True
        run.font.size = Pt(18)
        _set_chinese_font(run, "宋体")
        notice = document.add_paragraph(_NOTICE)
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._format_paragraph(notice, 9)
        if template_key == "id-loss-statement-v1":
            self._field_table(
                document,
                fields,
                (
                    ("申请人姓名", "applicant_name"),
                    ("联系电话", "contact_phone"),
                    ("遗失日期", "verification.loss_date"),
                    ("遗失地点", "verification.loss_place"),
                ),
            )
            self._body(
                document,
                "本人现就居民身份证遗失情况作如下说明：\n"
                + fields.get("verification.loss_description", "（请填写遗失经过）")
                + "\n本人确认以上内容为合成演示信息，并知悉本模板不作为正式申报材料。",
            )
            self._signature(document, "说明人")
        elif template_key in {
            "bl-agent-authorization-v1",
            "ess-operator-authorization-v1",
            "lc-operator-authorization-v1",
        }:
            agent_prefix = (
                "application"
                if template_key == "bl-agent-authorization-v1"
                else "operator"
            )
            self._field_table(
                document,
                fields,
                (
                    ("委托单位/经营主体", "business.name"),
                    ("统一社会信用代码", "business.unified_social_credit_code"),
                    ("法定代表人/经营者", "applicant_name"),
                    ("联系电话", "contact_phone"),
                    ("受托人", f"{agent_prefix}.agent_name" if agent_prefix == "application" else "operator.name"),
                    ("受托人演示证件标识", f"{agent_prefix}.agent_demo_id" if agent_prefix == "application" else "operator.demo_id"),
                    ("授权截止日期", f"{agent_prefix}.authorization_end_date"),
                ),
            )
            scope = fields.get(f"{agent_prefix}.authorization_scope", "（请填写授权事项）")
            self._body(
                document,
                f"现委托上述受托人办理与“{title_text.replace('（演示模板）', '')}”有关的演示事项。\n授权范围：{scope}\n受托人不得转委托，本授权不产生任何真实政务或法律效力。",
            )
            self._signature(document, "委托人")
        elif template_key == "ess-branch-statement-v1":
            self._field_table(
                document,
                fields,
                (
                    ("分支机构名称", "business.name"),
                    ("分支机构统一社会信用代码", "business.unified_social_credit_code"),
                    ("上级主体名称", "business.parent_name"),
                    ("上级主体统一社会信用代码", "business.parent_unified_social_credit_code"),
                    ("经办人", "applicant_name"),
                    ("联系电话", "contact_phone"),
                ),
            )
            self._body(document, "分支机构关系说明：\n" + fields.get("business.branch_relationship_description", "（请填写隶属关系及说明）"))
            self._signature(document, "说明单位")
        elif template_key == "lc-collective-statement-v1":
            self._field_table(
                document,
                fields,
                (
                    ("单位名称", "business.name"),
                    ("统一社会信用代码", "business.unified_social_credit_code"),
                    ("集体合同覆盖范围", "contract.collective_scope"),
                    ("涉及职工人数", "contract.employee_count"),
                    ("生效日期", "contract.effective_date"),
                    ("终止日期", "contract.end_date"),
                    ("经办人", "applicant_name"),
                    ("联系电话", "contact_phone"),
                ),
            )
            self._body(document, "本说明仅归纳合成演示集体合同的基本范围，不替代集体合同正文、职工代表程序或任何法定证明。")
            self._signature(document, "说明单位")
        elif template_key == "hf-personnel-summary-v1":
            self._field_table(
                document,
                fields,
                (
                    ("单位名称", "business.name"),
                    ("统一社会信用代码", "business.unified_social_credit_code"),
                    ("经办人", "applicant_name"),
                    ("联系电话", "contact_phone"),
                ),
            )
            rows = self._personnel_rows(fields.get("fund.demo_personnel", ""))
            table = document.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            for cell, label in zip(
                table.rows[0].cells,
                ("序号", "姓名（合成演示）", "缴存基数", "缴存比例", "备注"),
                strict=True,
            ):
                cell.text = label
            for index, item in enumerate(rows or [{}], start=1):
                cells = table.add_row().cells
                values = (
                    str(index),
                    str(item.get("name", "")),
                    str(item.get("base", fields.get("fund.contribution_base", ""))),
                    str(item.get("ratio", fields.get("fund.contribution_ratio", ""))),
                    str(item.get("remark", "")),
                )
                for cell, value in zip(cells, values, strict=True):
                    cell.text = value
            self._format_table(table)
        elif template_key == "hf-debit-authorization-v1":
            self._field_table(
                document,
                fields,
                (
                    ("授权单位", "business.name"),
                    ("统一社会信用代码", "business.unified_social_credit_code"),
                    ("演示开户行", "fund.demo_bank_name"),
                    ("演示账户名称", "fund.demo_account_name"),
                    ("演示账号", "fund.demo_account_number"),
                    ("经办人", "applicant_name"),
                    ("联系电话", "contact_phone"),
                ),
            )
            self._body(document, "本单位授权在项目演示环境中展示住房公积金委托扣款流程。该授权不连接银行、不触发扣款，也不具有真实法律效力。")
            self._signature(document, "授权单位")
        else:
            raise MaterialTemplateError("template_layout_unsupported")
        return document

    @staticmethod
    def _field_table(document: Any, fields: dict[str, str], specs: tuple[tuple[str, str], ...]) -> None:
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, key in specs:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = fields.get(key, "")
        DeterministicDocxMaterialRenderer._format_table(table)

    @staticmethod
    def _format_table(table: Any) -> None:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    DeterministicDocxMaterialRenderer._format_paragraph(paragraph, 10.5)

    @staticmethod
    def _format_paragraph(paragraph: Any, size: float = 10.5) -> None:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            _set_chinese_font(run, "宋体")

    @staticmethod
    def _body(document: Any, text: str) -> None:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.line_spacing = 1.5
        DeterministicDocxMaterialRenderer._format_paragraph(paragraph)

    @staticmethod
    def _signature(document: Any, label: str) -> None:
        paragraph = document.add_paragraph(f"\n{label}（签字/盖章）：________________\n日期：____年__月__日")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        DeterministicDocxMaterialRenderer._format_paragraph(paragraph)

    @staticmethod
    def _personnel_rows(value: str) -> list[dict[str, Any]]:
        if not value.strip():
            return []
        try:
            raw = json.loads(value)
        except Exception:
            return []
        if not isinstance(raw, list) or len(raw) > 50:
            return []
        return [item for item in raw if isinstance(item, dict)][:50]

    @staticmethod
    def _replace_placeholders(
        document: Any, fields: dict[str, str]
    ) -> set[str]:
        matched: set[str] = set()
        containers = list(document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    containers.extend(cell.paragraphs)
        for paragraph in containers:
            original = paragraph.text
            updated = original
            for key, value in fields.items():
                before = updated
                updated = updated.replace("{{" + key + "}}", value)
                updated = updated.replace("[[" + key + "]]", value)
                if updated != before:
                    matched.add(key)
            if updated != original:
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = updated
                else:
                    paragraph.add_run(updated)
        return matched

    @staticmethod
    def _fill_labeled_blanks(
        document: Any, fields: dict[str, str], matched: set[str]
    ) -> None:
        aliases = {
            "business.name": ("甲方名称", "用人单位名称", "单位名称"),
            "business.unified_social_credit_code": ("统一社会信用代码",),
            "contract.employee_name": ("乙方姓名", "劳动者姓名"),
            "contract.employee_demo_id": ("身份证号", "证件号码"),
            "contract.effective_date": ("合同开始日期", "合同期限自"),
            "contract.end_date": ("合同终止日期", "合同期限至"),
            "contract.work_role": ("乙方同意在", "工作岗位", "岗位"),
            "contract.work_location": ("工作地点",),
            "contract.demo_salary": ("工资标准", "劳动报酬"),
        }
        paragraphs = list(document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        used_paragraphs: set[int] = set()
        for key, value in fields.items():
            if key in matched:
                continue
            for alias in aliases.get(key, ()):
                for paragraph in paragraphs:
                    label_text = paragraph.text.strip()
                    if (
                        id(paragraph) in used_paragraphs
                        or alias not in label_text
                        or len(label_text) > 100
                        or label_text.index(alias) > 20
                    ):
                        continue
                    if value in paragraph.text:
                        matched.add(key)
                        break
                    if not DeterministicDocxMaterialRenderer._replace_labeled_blank(
                        paragraph, value
                    ):
                        # A heading or narrative sentence can contain the same
                        # label. Only a reviewed reserved blank is a stable fill
                        # target; otherwise keep searching and leave unmatched
                        # fields empty.
                        continue
                    used_paragraphs.add(id(paragraph))
                    matched.add(key)
                    break
                if key in matched:
                    break

    @staticmethod
    def _replace_labeled_blank(paragraph: Any, value: str) -> bool:
        """Replace a template's first underlined blank without widening the line.

        Reviewed source forms commonly model a field as many whitespace-only runs.
        Appending a value after those runs pushes it to the next line, so consume the
        reserved blank and reuse its formatting instead.
        """

        runs = list(paragraph.runs)
        start: int | None = None
        for index, run in enumerate(runs):
            if run.text and not run.text.strip() and bool(run.underline):
                start = index
                break
        if start is None:
            return False
        target = runs[start]
        target.text = value
        target.underline = True
        _set_chinese_font(target, "宋体")
        for run in runs[start + 1 :]:
            if run.text and run.text.strip():
                break
            run.text = ""
        return True

    @staticmethod
    def _add_notice_footer(document: Any) -> None:
        for section in document.sections:
            paragraph = section.footer.paragraphs[0]
            paragraph.text = _NOTICE
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8)
                _set_chinese_font(run, "宋体")

    @staticmethod
    def _scrub_properties(document: Any) -> None:
        props = document.core_properties
        props.author = ""
        props.last_modified_by = ""
        props.comments = ""
        props.keywords = ""
        props.subject = ""
        props.category = ""


def _set_chinese_font(run: Any, name: str) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _lookup(values: dict[str, Any], dotted: str) -> Any:
    current: Any = values
    for part in dotted.split("."):
        if not isinstance(current, dict) or part not in current:
            return None
        current = current[part]
    return current


def _bounded_snapshot(
    values: dict[str, Any], allowed_fields: tuple[str, ...]
) -> dict[str, str]:
    result: dict[str, str] = {}
    for field in allowed_fields:
        value = _lookup(values, field)
        if isinstance(value, (str, int, float, bool)):
            text = str(value).strip()
            if text:
                result[field] = text[:500]
        elif isinstance(value, (list, dict)):
            encoded = json.dumps(value, ensure_ascii=False, separators=(",", ":"))
            if len(encoded) <= 4000:
                result[field] = encoded
    return result


def _decode_analysis(raw: Any, allowed_fields: tuple[str, ...]) -> TemplateAnalysis:
    if not isinstance(raw, dict):
        raise MaterialTemplateError("model_schema_invalid")
    fields = raw.get("fields", {})
    normalized: dict[str, str] = {}
    if isinstance(fields, dict):
        # Read by allowlist instead of iterating model-authored keys. Unknown
        # keys and non-scalar values are intentionally ignored.
        for key in allowed_fields:
            value = fields.get(key)
            if not isinstance(value, (str, int, float, bool)):
                continue
            text = re.sub(r"\s+", " ", str(value)).strip()
            if text:
                normalized[key] = text[:4000]
    warnings_raw = raw.get("warnings", [])
    warnings: list[str] = []
    if isinstance(warnings_raw, list):
        for value in warnings_raw:
            if len(warnings) >= 8:
                break
            if not isinstance(value, str):
                continue
            text = re.sub(r"\s+", " ", value).strip()
            if text and len(text) <= 200:
                warnings.append(text)
    return TemplateAnalysis(normalized, tuple(warnings))
