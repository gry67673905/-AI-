from __future__ import annotations

import asyncio
import hashlib
import io
import json
from datetime import datetime, timedelta, timezone
from pathlib import Path
from types import SimpleNamespace
from uuid import uuid4

import httpx
import pytest
from docx import Document
from pydantic import ValidationError

from app.application.dtos import Principal
from app.application.material_documents import (
    LeasedMaterialDocumentJob,
    MaterialDocumentCoordinator,
    MaterialDocumentWorker,
    TemplateAnalysis,
)
from app.domain.enums import (
    ApplicantType,
    ApplicationStatus,
    MaterialTemplateMode,
    Role,
    ServiceStatus,
)
from app.config import Settings
from app.infrastructure.material_documents import (
    DashScopeMaterialTemplateAnalyzer,
    DeterministicDocxMaterialRenderer,
    MaterialTemplateError,
    MaterialTemplatePack,
    MockMaterialTemplateAnalyzer,
    _decode_analysis,
    validate_docx_package,
)
from app.infrastructure.material_template_catalog import (
    immutable_material_template_objects,
)
from app.infrastructure.object_store import InMemoryObjectStore, MinioObjectStore
from app.infrastructure.records import (
    ApplicationRecord,
    GovernmentServiceRecord,
    MaterialDocumentJobRecord,
    MaterialTemplateRecord,
)
from app.infrastructure.repositories import (
    BusinessRepository,
    _enforce_material_document_job_limits,
)
from app.infrastructure.runtime import BusinessRuntime
from app.ops.material_template_check import check_material_template_pack
from app.errors import ConflictError, ResourceNotFound, TooManyRequests


def _docx(text: str = "姓名：{{applicant_name}}") -> bytes:
    document = Document()
    document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _citizen() -> Principal:
    return Principal(
        account_id=uuid4(),
        username="citizen",
        display_name="演示群众",
        role=Role.CITIZEN,
        applicant_type=ApplicantType.INDIVIDUAL,
        token_version=0,
    )


def test_reviewed_template_pack_is_hash_checked_and_bounded() -> None:
    pack = MaterialTemplatePack(
        "resources/material_templates/v1/manifest.json", "material-templates/v1"
    )
    templates = pack.load()
    assert len(templates) == 18
    generated = [
        item
        for item in templates
        if item.mode is not MaterialTemplateMode.NOT_GENERATABLE
    ]
    assert generated
    assert all(item.source_bytes and item.source_sha256 for item in generated)
    assert all(
        hashlib.sha256(item.source_bytes or b"").hexdigest()
        == item.source_sha256
        for item in generated
    )


def test_template_object_keys_are_versioned_and_content_addressed() -> None:
    templates = immutable_material_template_objects(
        MaterialTemplatePack(
            "resources/material_templates/v1/manifest.json", "material-templates/v1"
        ).load()
    )
    generated = [item for item in templates if item.source_sha256 is not None]
    assert generated
    assert all(
        item.source_object_key
        == (
            f"material-templates/v1/{item.template_key}/v1/"
            f"{item.source_sha256}.docx"
        )
        for item in generated
    )


def test_material_job_schema_contains_immutable_generation_snapshot() -> None:
    columns = MaterialDocumentJobRecord.__table__.c
    assert {
        "template_key_snapshot",
        "template_version_snapshot",
        "template_title_snapshot",
        "template_mode_snapshot",
        "allowed_fields_snapshot",
        "source_object_key_snapshot",
        "source_sha256_snapshot",
        "model_name_snapshot",
        "release_lane",
    }.issubset(columns.keys())


@pytest.mark.parametrize(
    ("left", "right"),
    [
        ("materials_bucket", "knowledge_bucket"),
        ("materials_bucket", "material_templates_bucket"),
        ("materials_bucket", "generated_documents_bucket"),
        ("knowledge_bucket", "material_templates_bucket"),
        ("knowledge_bucket", "generated_documents_bucket"),
        ("material_templates_bucket", "generated_documents_bucket"),
    ],
)
def test_settings_require_four_distinct_object_store_buckets(
    left: str, right: str
) -> None:
    values = {
        "materials_bucket": "materials",
        "knowledge_bucket": "knowledge",
        "material_templates_bucket": "templates",
        "generated_documents_bucket": "generated",
    }
    values[left] = "COLLISION"
    values[right] = "collision"
    with pytest.raises(ValidationError, match="dedicated bucket distinct"):
        Settings(_env_file=None, **values)  # type: ignore[arg-type]


def test_cloud_preflight_requires_bucket_isolation_and_global_daily_cap() -> None:
    script = (
        Path(__file__).parents[2] / "deploy" / "scripts" / "lib.sh"
    ).read_text(encoding="utf-8")
    for key in (
        "MATERIALS_BUCKET",
        "KNOWLEDGE_BUCKET",
        "MATERIAL_TEMPLATES_BUCKET",
        "GENERATED_DOCUMENTS_BUCKET",
        "MATERIAL_DOCUMENT_GLOBAL_DAILY_LIMIT",
    ):
        assert f"cloud_env_value {key}" in script
    assert "four distinct MinIO buckets" in script


def test_backup_keeps_template_sources_and_omits_disposable_generated_files() -> None:
    root = Path(__file__).parents[2]
    entrypoint = (root / "deploy" / "container-secret-entrypoint.sh").read_text(
        encoding="utf-8"
    )
    backup = (root / "deploy" / "scripts" / "backup.sh").read_text(
        encoding="utf-8"
    )
    backup_case = entrypoint.split("minio-backup)", 1)[1].split(";;", 1)[0]
    assert "MATERIAL_TEMPLATES_BUCKET" in backup_case
    assert "/backup/material-templates" in backup_case
    assert "GENERATED_DOCUMENTS_BUCKET" not in backup_case
    assert "generated_documents_backup=omitted_disposable_24h" in backup


@pytest.mark.asyncio
async def test_old_draft_lists_templates_from_its_pinned_service_version() -> None:
    owner_id = uuid4()
    application_id = uuid4()
    service_id = uuid4()
    old_version_id = uuid4()
    newly_published_version_id = uuid4()
    template_id = uuid4()
    captured_statement: object | None = None

    application = SimpleNamespace(
        id=application_id,
        applicant_id=owner_id,
        service_id=service_id,
        service_version_id=old_version_id,
    )
    service = SimpleNamespace(
        id=service_id,
        status=ServiceStatus.PUBLISHED.value,
        current_version_id=newly_published_version_id,
    )
    requirement = SimpleNamespace(
        code="old-requirement",
        name="旧版材料",
        order_index=0,
    )
    template = SimpleNamespace(
        id=template_id,
        active=True,
        mode=MaterialTemplateMode.SOURCE_EDITABLE.value,
        source_object_key="material-templates/v1/old/v1/digest.docx",
        source_sha256="a" * 64,
        title="旧版模板",
        notice="旧草稿固定模板",
        template_key="old-template-v1",
    )

    class Rows:
        @staticmethod
        def all() -> list[tuple[object, object]]:
            return [(requirement, template)]

    class Session:
        async def get(self, model: object, identity: object) -> object | None:
            if model is ApplicationRecord and identity == application_id:
                return application
            if model is GovernmentServiceRecord and identity == service_id:
                return service
            return None

        async def execute(self, statement: object) -> Rows:
            nonlocal captured_statement
            captured_statement = statement
            return Rows()

    class SessionContext:
        async def __aenter__(self) -> Session:
            return Session()

        async def __aexit__(self, *_args: object) -> None:
            return None

    repository = BusinessRepository(lambda: SessionContext())  # type: ignore[arg-type]
    items = await repository.list_material_template_options(application_id, owner_id)
    assert items[0]["template_id"] == template_id
    assert captured_statement is not None
    parameters = captured_statement.compile().params  # type: ignore[union-attr]
    assert old_version_id in parameters.values()
    assert newly_published_version_id not in parameters.values()
    with pytest.raises(ResourceNotFound):
        await repository.list_material_template_options(application_id, uuid4())


def _catalog_template(**overrides: object) -> SimpleNamespace:
    values: dict[str, object] = {
        "template_key": "test-template-v1",
        "service_code": "DEMO-TEST-001",
        "requirement_code": "test-1",
        "title": "测试模板",
        "mode": MaterialTemplateMode.SOURCE_EDITABLE,
        "source_object_key": "material-templates/v1/test/v1/" + "a" * 64 + ".docx",
        "source_sha256": "a" * 64,
        "allowed_fields": ("applicant_name",),
        "notice": "演示模板。",
    }
    values.update(overrides)
    return SimpleNamespace(**values)


class _SeedSession:
    def __init__(self, scalar_results: list[object | None]) -> None:
        self.scalar_results = iter(scalar_results)
        self.executed: list[object] = []
        self.added: list[object] = []

    async def execute(self, statement: object, *_args: object) -> None:
        self.executed.append(statement)

    async def scalar(self, _statement: object) -> object | None:
        return next(self.scalar_results)

    def add(self, record: object) -> None:
        self.added.append(record)


async def _run_seed(
    repository: BusinessRepository,
    session: _SeedSession,
    templates: tuple[object, ...],
) -> None:
    token = repository._uow_session.set(session)  # type: ignore[arg-type]
    try:
        await repository.seed_material_templates(templates)
    finally:
        repository._uow_session.reset(token)


@pytest.mark.asyncio
async def test_template_seed_fails_on_missing_service_reference() -> None:
    repository = BusinessRepository(object())  # type: ignore[arg-type]
    session = _SeedSession([None])
    with pytest.raises(RuntimeError, match="material_template_service_missing"):
        await _run_seed(repository, session, (_catalog_template(),))
    assert len(session.executed) == 1


@pytest.mark.asyncio
async def test_template_seed_fails_on_missing_requirement_reference() -> None:
    repository = BusinessRepository(object())  # type: ignore[arg-type]
    session = _SeedSession([SimpleNamespace(current_version_id=uuid4()), None])
    with pytest.raises(RuntimeError, match="material_template_requirement_missing"):
        await _run_seed(repository, session, (_catalog_template(),))
    assert len(session.executed) == 1


@pytest.mark.asyncio
async def test_template_seed_deactivates_removed_keys_without_touching_present_versions() -> None:
    repository = BusinessRepository(object())  # type: ignore[arg-type]
    requirement_id = uuid4()
    template = _catalog_template()
    existing = SimpleNamespace(
        material_requirement_id=requirement_id,
        service_code=template.service_code,
        requirement_code=template.requirement_code,
        title=template.title,
        mode=template.mode.value,
        source_object_key=template.source_object_key,
        source_sha256=template.source_sha256,
        allowed_fields=list(template.allowed_fields),
        notice=template.notice,
        active=True,
        version=1,
    )
    session = _SeedSession(
        [
            SimpleNamespace(current_version_id=uuid4()),
            SimpleNamespace(id=requirement_id),
            existing,
        ]
    )
    await _run_seed(repository, session, (template,))

    updates = session.executed[1:]
    assert len(updates) == 1
    deactivate_sql = str(updates[0])
    assert "template_key NOT IN" in deactivate_sql
    assert "material_requirement_id" not in deactivate_sql
    assert template.template_key in str(updates[0].compile().params)
    assert existing.active is True


@pytest.mark.asyncio
async def test_template_reintroduction_activates_only_current_version_row() -> None:
    repository = BusinessRepository(object())  # type: ignore[arg-type]
    requirement_id = uuid4()
    template = _catalog_template()
    recalled_historical = SimpleNamespace(active=False)
    session = _SeedSession(
        [
            SimpleNamespace(current_version_id=uuid4()),
            SimpleNamespace(id=requirement_id),
            None,
        ]
    )
    await _run_seed(repository, session, (template,))

    assert len(session.executed[1:]) == 1
    assert "template_key NOT IN" in str(session.executed[1])
    assert recalled_historical.active is False
    assert len(session.added) == 1
    current = session.added[0]
    assert isinstance(current, MaterialTemplateRecord)
    assert current.material_requirement_id == requirement_id
    assert current.active is True


@pytest.mark.asyncio
async def test_template_seed_rejects_in_place_semantic_or_source_change() -> None:
    repository = BusinessRepository(object())  # type: ignore[arg-type]
    requirement_id = uuid4()
    template = _catalog_template(title="新标题")
    existing = SimpleNamespace(
        material_requirement_id=requirement_id,
        service_code=template.service_code,
        requirement_code=template.requirement_code,
        title="旧标题",
        mode=template.mode.value,
        source_object_key=template.source_object_key,
        source_sha256=template.source_sha256,
        allowed_fields=list(template.allowed_fields),
        notice=template.notice,
        active=True,
        version=1,
    )
    session = _SeedSession(
        [
            SimpleNamespace(current_version_id=uuid4()),
            SimpleNamespace(id=requirement_id),
            existing,
        ]
    )
    with pytest.raises(RuntimeError, match="material_template_immutable_conflict"):
        await _run_seed(repository, session, (template,))
    assert len(session.executed) == 1
    assert existing.title == "旧标题"


def _manifest_item(**overrides: object) -> dict[str, object]:
    item: dict[str, object] = {
        "template_id": "test-template-v1",
        "service_code": "DEMO-TEST-001",
        "requirement_code": "test-1",
        "title": "测试模板",
        "mode": "SOURCE_EDITABLE",
        "packaged_path": "docx/test.docx",
        "packaged_sha256": "0" * 64,
        "allowed_fields": ["applicant_name"],
        "notice": "演示模板。",
    }
    item.update(overrides)
    return item


def _write_manifest(path: Path, templates: list[dict[str, object]]) -> None:
    manifest = path / "manifest.json"
    manifest.write_text(
        json.dumps({"version": 1, "templates": templates}, ensure_ascii=False),
        encoding="utf-8",
    )


def test_template_pack_rejects_missing_and_empty_manifest(tmp_path: Path) -> None:
    with pytest.raises(MaterialTemplateError, match="template_manifest_missing"):
        MaterialTemplatePack(str(tmp_path / "missing.json"), "pack/v1").load()

    _write_manifest(tmp_path, [])
    with pytest.raises(MaterialTemplateError, match="template_manifest_items"):
        MaterialTemplatePack(str(tmp_path / "manifest.json"), "pack/v1").load()


@pytest.mark.parametrize(
    ("override", "expected"),
    [
        ({"packaged_path": None}, "template_manifest_source"),
        ({"packaged_path": "docx/missing.docx"}, "template_manifest_source_missing"),
    ],
)
def test_generatable_template_requires_packaged_source(
    tmp_path: Path, override: dict[str, object], expected: str
) -> None:
    _write_manifest(tmp_path, [_manifest_item(**override)])
    with pytest.raises(MaterialTemplateError, match=expected):
        MaterialTemplatePack(str(tmp_path / "manifest.json"), "pack/v1").load()


@pytest.mark.parametrize("declared_sha", [None, "not-a-sha", "0" * 64])
def test_generatable_template_requires_matching_packaged_sha256(
    tmp_path: Path, declared_sha: str | None
) -> None:
    content = _docx()
    docx_dir = tmp_path / "docx"
    docx_dir.mkdir()
    (docx_dir / "test.docx").write_bytes(content)
    _write_manifest(
        tmp_path,
        [_manifest_item(packaged_sha256=declared_sha)],
    )
    with pytest.raises(MaterialTemplateError, match="template_manifest_sha256"):
        MaterialTemplatePack(str(tmp_path / "manifest.json"), "pack/v1").load()


def test_disabled_runtime_does_not_load_material_template_manifest(
    tmp_path: Path,
) -> None:
    settings = Settings(
        ENVIRONMENT="local",
        ENABLE_DEMO_PROVIDERS=False,
        MATERIAL_DOCUMENTS_ENABLED=False,
        MATERIAL_TEMPLATE_MANIFEST_PATH=str(tmp_path / "missing.json"),
    )
    runtime = BusinessRuntime(settings, None, None, None, None, None, None)
    assert runtime._packed_material_templates == ()


def test_nonpaid_pack_check_requires_enabled_nonempty_sources() -> None:
    settings = Settings(MATERIAL_DOCUMENTS_ENABLED=True)
    assert check_material_template_pack(settings) == (18, 9)

    disabled = Settings(MATERIAL_DOCUMENTS_ENABLED=False)
    with pytest.raises(MaterialTemplateError, match="material_documents_disabled"):
        check_material_template_pack(disabled)


def test_nonpaid_pack_check_rejects_catalog_without_generatable_sources(
    tmp_path: Path,
) -> None:
    _write_manifest(
        tmp_path,
        [
            _manifest_item(
                mode="NOT_GENERATABLE",
                packaged_path=None,
                packaged_sha256=None,
            )
        ],
    )
    settings = Settings(
        MATERIAL_DOCUMENTS_ENABLED=True,
        MATERIAL_TEMPLATE_MANIFEST_PATH=str(tmp_path / "manifest.json"),
    )
    with pytest.raises(MaterialTemplateError, match="template_pack_sources_empty"):
        check_material_template_pack(settings)


@pytest.mark.asyncio
async def test_deterministic_renderer_fills_only_markers_and_scrubs_metadata() -> None:
    source = _docx()
    renderer = DeterministicDocxMaterialRenderer()
    result = await renderer.render(
        mode=MaterialTemplateMode.SOURCE_EDITABLE,
        template_key="lc-labor-contract-v1",
        template_title="丢失情况说明",
        source_docx=source,
        analysis=TemplateAnalysis({"applicant_name": "演示甲"}),
    )
    validate_docx_package(result)
    reopened = Document(io.BytesIO(result))
    assert "姓名：演示甲" in "\n".join(item.text for item in reopened.paragraphs)
    assert reopened.core_properties.author == ""
    assert "演示模板" in reopened.sections[0].footer.paragraphs[0].text


@pytest.mark.asyncio
async def test_source_editable_renderer_consumes_reserved_underlined_blank() -> None:
    source = Document()
    source.add_paragraph("工作地点")
    paragraph = source.add_paragraph()
    paragraph.add_run("甲方名称：")
    blank = paragraph.add_run(" " * 48)
    blank.underline = True
    paragraph.add_run(" " * 24)
    buffer = io.BytesIO()
    source.save(buffer)

    result = await DeterministicDocxMaterialRenderer().render(
        mode=MaterialTemplateMode.SOURCE_EDITABLE,
        template_key="lc-labor-contract-v1",
        template_title="劳动合同演示件",
        source_docx=buffer.getvalue(),
        analysis=TemplateAnalysis({"business.name": "成都未来演示科技有限公司"}),
    )

    reopened = Document(io.BytesIO(result))
    rendered = reopened.paragraphs[0].text
    assert rendered == "工作地点"
    assert reopened.paragraphs[1].text == "甲方名称：成都未来演示科技有限公司"


@pytest.mark.asyncio
async def test_source_editable_renderer_skips_alias_heading_for_fillable_line() -> None:
    source = Document()
    source.add_paragraph("二、工作内容和工作地点")
    paragraph = source.add_paragraph()
    paragraph.add_run("双方确认工作地点为：")
    blank = paragraph.add_run(" " * 24)
    blank.underline = True
    paragraph.add_run("。")
    buffer = io.BytesIO()
    source.save(buffer)

    result = await DeterministicDocxMaterialRenderer().render(
        mode=MaterialTemplateMode.SOURCE_EDITABLE,
        template_key="lc-labor-contract-v1",
        template_title="劳动合同演示件",
        source_docx=buffer.getvalue(),
        analysis=TemplateAnalysis({"contract.work_location": "成都市演示办公区"}),
    )

    reopened = Document(io.BytesIO(result))
    assert reopened.paragraphs[0].text == "二、工作内容和工作地点"
    assert reopened.paragraphs[1].text == "双方确认工作地点为：成都市演示办公区。"


@pytest.mark.asyncio
async def test_reconstructed_personnel_summary_uses_controlled_labels_and_rows() -> None:
    renderer = DeterministicDocxMaterialRenderer()
    result = await renderer.render(
        mode=MaterialTemplateMode.VISUAL_RECONSTRUCT,
        template_key="hf-personnel-summary-v1",
        template_title="缴存人员汇总表（演示模板）",
        source_docx=_docx("参考版式"),
        analysis=TemplateAnalysis(
            {
                "business.name": "演示单位",
                "fund.demo_personnel": '[{"name":"演示员工甲","base":"5000","ratio":"8%"}]',
            }
        ),
    )
    document = Document(io.BytesIO(result))
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "姓名（合成演示）" in table_text
    assert "演示员工甲" in table_text
    assert "fund.demo_personnel" not in table_text


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "template_key",
    [
        "id-loss-statement-v1",
        "bl-agent-authorization-v1",
        "ess-operator-authorization-v1",
        "ess-branch-statement-v1",
        "lc-operator-authorization-v1",
        "lc-collective-statement-v1",
        "hf-personnel-summary-v1",
        "hf-debit-authorization-v1",
    ],
)
async def test_every_visual_template_uses_a_supported_controlled_layout(
    template_key: str,
) -> None:
    renderer = DeterministicDocxMaterialRenderer()
    result = await renderer.render(
        mode=MaterialTemplateMode.VISUAL_RECONSTRUCT,
        template_key=template_key,
        template_title="受控演示模板",
        source_docx=_docx("仅用于视觉版式参考"),
        analysis=TemplateAnalysis({}),
    )
    validate_docx_package(result)
    document = Document(io.BytesIO(result))
    assert document.tables
    assert "演示模板" in document.sections[0].footer.paragraphs[0].text


@pytest.mark.asyncio
async def test_dashscope_analyzer_makes_exactly_one_structured_request() -> None:
    calls = 0

    async def handler(request: httpx.Request) -> httpx.Response:
        nonlocal calls
        calls += 1
        body = __import__("json").loads(request.content)
        assert body["response_format"] == {"type": "json_object"}
        assert body["enable_thinking"] is False
        return httpx.Response(
            200,
            json={
                "choices": [
                    {
                        "message": {
                            "content": '{"fields":{"applicant_name":"演示甲"},"warnings":[]}'
                        }
                    }
                ]
            },
        )

    client = httpx.AsyncClient(transport=httpx.MockTransport(handler))
    analyzer = DashScopeMaterialTemplateAnalyzer(
        "test-key",
        model_name="qwen3-vl-flash-2026-01-22",
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
        timeout_seconds=90,
        client=client,
    )
    result = await analyzer.analyze(
        page_images=(b"\xff\xd8\xffdemo",),
        template_title="演示模板",
        allowed_fields=("applicant_name",),
        form_snapshot={"applicant_name": "演示甲"},
        request_text=None,
    )
    assert result.fields == {"applicant_name": "演示甲"}
    assert calls == 1
    await client.aclose()


@pytest.mark.asyncio
async def test_dashscope_ignores_extra_metadata_unknown_fields_and_bad_warnings() -> None:
    calls = 0

    async def handler(_request: httpx.Request) -> httpx.Response:
        nonlocal calls
        calls += 1
        content = {
            "fields": {
                "applicant_name": "  演示   甲  ",
                "applicant_age": 26,
                "structured": {"must": "be ignored"},
                "unknown_secret": "must be ignored",
            },
            "warnings": [
                None,
                123,
                "",
                "x" * 201,
                " 有效提示 0 ",
                *[f"有效提示 {index}" for index in range(1, 12)],
            ],
            "request_id": "provider-metadata",
            "usage": {"tokens": 123},
        }
        return httpx.Response(
            200,
            json={
                "choices": [
                    {"message": {"content": json.dumps(content, ensure_ascii=False)}}
                ]
            },
        )

    client = httpx.AsyncClient(transport=httpx.MockTransport(handler))
    analyzer = DashScopeMaterialTemplateAnalyzer(
        "test-key",
        model_name="qwen3-vl-flash-2026-01-22",
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
        timeout_seconds=90,
        client=client,
    )
    result = await analyzer.analyze(
        page_images=(b"\xff\xd8\xffdemo",),
        template_title="演示模板",
        allowed_fields=("applicant_name", "applicant_age", "structured"),
        form_snapshot={},
        request_text=None,
    )
    assert result.fields == {
        "applicant_name": "演示 甲",
        "applicant_age": "26",
    }
    assert result.warnings == tuple(f"有效提示 {index}" for index in range(8))
    assert calls == 1
    await client.aclose()


@pytest.mark.asyncio
async def test_dashscope_invalid_json_fails_once_without_retry_or_fallback() -> None:
    calls = 0

    async def handler(_request: httpx.Request) -> httpx.Response:
        nonlocal calls
        calls += 1
        return httpx.Response(
            200,
            json={
                "choices": [
                    {"message": {"content": "not-json provider detail"}}
                ]
            },
        )

    client = httpx.AsyncClient(transport=httpx.MockTransport(handler))
    analyzer = DashScopeMaterialTemplateAnalyzer(
        "test-key",
        model_name="qwen3-vl-flash-2026-01-22",
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
        timeout_seconds=90,
        client=client,
    )
    with pytest.raises(MaterialTemplateError) as caught:
        await analyzer.analyze(
            page_images=(b"\xff\xd8\xffdemo",),
            template_title="演示模板",
            allowed_fields=("applicant_name",),
            form_snapshot={},
            request_text=None,
        )
    assert caught.value.code == "model_json_invalid"
    assert calls == 1
    await client.aclose()


@pytest.mark.asyncio
async def test_dashscope_http_error_maps_to_static_code_without_retry() -> None:
    calls = 0

    async def handler(_request: httpx.Request) -> httpx.Response:
        nonlocal calls
        calls += 1
        return httpx.Response(500, text="provider secret diagnostic")

    client = httpx.AsyncClient(transport=httpx.MockTransport(handler))
    analyzer = DashScopeMaterialTemplateAnalyzer(
        "test-key",
        model_name="qwen3-vl-flash-2026-01-22",
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
        timeout_seconds=90,
        client=client,
    )
    with pytest.raises(MaterialTemplateError) as caught:
        await analyzer.analyze(
            page_images=(b"\xff\xd8\xffdemo",),
            template_title="演示模板",
            allowed_fields=("applicant_name",),
            form_snapshot={},
            request_text=None,
        )
    assert caught.value.code == "model_request_failed"
    assert str(caught.value) == "model_request_failed"
    assert calls == 1
    await client.aclose()


def test_material_template_error_exposes_only_allowlisted_static_code() -> None:
    assert MaterialTemplateError("model_json_invalid").code == "model_json_invalid"
    unsafe = MaterialTemplateError("provider said bearer secret-value")
    assert unsafe.code == "material_template_error"
    assert str(unsafe) == "material_template_error"


def test_analysis_ignores_malformed_optional_containers() -> None:
    result = _decode_analysis(
        {
            "fields": ["not", "an", "object"],
            "warnings": {"not": "a list"},
            "metadata": True,
        },
        ("applicant_name",),
    )
    assert result == TemplateAnalysis(fields={}, warnings=())


class _WorkerRepository:
    def __init__(self, job: LeasedMaterialDocumentJob) -> None:
        self.job = job
        self.completed: dict[str, object] | None = None
        self.failed: str | None = None
        self.renewals = 0
        self.renew_success = True
        self.leased_lane: str | None = None

    async def lease_material_document_job(
        self, _lease_seconds: int, release_lane: str
    ) -> LeasedMaterialDocumentJob | None:
        self.leased_lane = release_lane
        job, self.job = self.job, None  # type: ignore[assignment]
        return job

    async def renew_material_document_job_lease(
        self, _generation_id: object, _lease_token: object, _lease_seconds: int
    ) -> bool:
        self.renewals += 1
        return self.renew_success

    async def complete_material_document_job(
        self,
        _generation_id: object,
        _lease_token: object,
        object_key: str,
        _filename: str,
        size_bytes: int,
        sha256: str,
        _model_name: str,
        _warnings: tuple[str, ...],
    ) -> bool:
        self.completed = {
            "object_key": object_key,
            "size_bytes": size_bytes,
            "sha256": sha256,
        }
        return True

    async def fail_material_document_job(
        self, _generation_id: object, _lease_token: object, error_code: str
    ) -> bool:
        self.failed = error_code
        return True

    async def expire_material_document_jobs(
        self, _now: datetime
    ) -> list[str]:
        return []


class _PageRenderer:
    async def render_pages(self, _source: bytes) -> tuple[bytes, ...]:
        return (b"jpeg-page",)


class _SlowPageRenderer:
    def __init__(self, delay: float) -> None:
        self.delay = delay

    async def render_pages(self, _source: bytes) -> tuple[bytes, ...]:
        await asyncio.sleep(self.delay)
        return (b"jpeg-page",)


class _FailingAnalyzer:
    model_name = MockMaterialTemplateAnalyzer.model_name

    def __init__(self, error: Exception) -> None:
        self.error = error
        self.calls = 0

    async def analyze(self, **_kwargs: object) -> TemplateAnalysis:
        self.calls += 1
        raise self.error

    async def close(self) -> None:
        return None


@pytest.mark.asyncio
async def test_worker_writes_private_docx_and_never_falls_back() -> None:
    source = _docx()
    owner = uuid4()
    job = LeasedMaterialDocumentJob(
        id=uuid4(),
        owner_account_id=owner,
        application_id=uuid4(),
        requirement_code="id-2",
        template_id=uuid4(),
        template_key="id-loss-statement-v1",
        template_title="丢失情况说明",
        template_mode=MaterialTemplateMode.SOURCE_EDITABLE,
        allowed_fields=("applicant_name",),
        source_object_key="material-templates/v1/source.docx",
        source_sha256=hashlib.sha256(source).hexdigest(),
        form_snapshot={"applicant_name": "演示甲"},
        request_text=None,
        lease_token=uuid4(),
    )
    repository = _WorkerRepository(job)
    store = InMemoryObjectStore()
    await store.put_bytes(
        store.material_templates_bucket,
        job.source_object_key or "",
        source,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    worker = MaterialDocumentWorker(
        repository,  # type: ignore[arg-type]
        store,
        MockMaterialTemplateAnalyzer(),
        _PageRenderer(),
        DeterministicDocxMaterialRenderer(),
        lease_seconds=180,
        release_lane="release-test",
    )
    assert await worker.run_once() is True
    assert repository.leased_lane == "release-test"
    assert repository.failed is None
    assert repository.completed is not None
    key = str(repository.completed["object_key"])
    result = await store.get_bytes(store.generated_documents_bucket, key)
    assert "演示甲" in "\n".join(
        item.text for item in Document(io.BytesIO(result)).paragraphs
    )


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("error", "expected_code"),
    [
        (MaterialTemplateError("model_json_invalid"), "MODEL_JSON_INVALID"),
        (RuntimeError("Bearer provider-secret response"), "UNEXPECTED_WORKER_ERROR"),
    ],
)
async def test_worker_persists_only_safe_static_failure_code(
    error: Exception, expected_code: str
) -> None:
    source = _docx()
    job = LeasedMaterialDocumentJob(
        id=uuid4(),
        owner_account_id=uuid4(),
        application_id=uuid4(),
        requirement_code="id-2",
        template_id=uuid4(),
        template_key="id-loss-statement-v1",
        template_title="丢失情况说明",
        template_mode=MaterialTemplateMode.SOURCE_EDITABLE,
        allowed_fields=("applicant_name",),
        source_object_key="material-templates/v1/source.docx",
        source_sha256=hashlib.sha256(source).hexdigest(),
        form_snapshot={},
        request_text=None,
        lease_token=uuid4(),
    )
    repository = _WorkerRepository(job)
    analyzer = _FailingAnalyzer(error)
    store = InMemoryObjectStore()
    await store.put_bytes(
        store.material_templates_bucket,
        job.source_object_key or "",
        source,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    worker = MaterialDocumentWorker(
        repository,  # type: ignore[arg-type]
        store,
        analyzer,  # type: ignore[arg-type]
        _PageRenderer(),
        DeterministicDocxMaterialRenderer(),
        lease_seconds=180,
    )
    assert await worker.run_once() is True
    assert analyzer.calls == 1
    assert repository.completed is None
    assert repository.failed == expected_code
    assert "SECRET" not in (repository.failed or "")


@pytest.mark.asyncio
async def test_worker_renews_lease_while_a_valid_job_is_slow() -> None:
    source = _docx()
    job = LeasedMaterialDocumentJob(
        id=uuid4(),
        owner_account_id=uuid4(),
        application_id=uuid4(),
        requirement_code="id-2",
        template_id=uuid4(),
        template_key="id-loss-statement-v1",
        template_title="丢失情况说明",
        template_mode=MaterialTemplateMode.SOURCE_EDITABLE,
        allowed_fields=("applicant_name",),
        source_object_key="material-templates/v1/source.docx",
        source_sha256=hashlib.sha256(source).hexdigest(),
        form_snapshot={"applicant_name": "演示甲"},
        request_text=None,
        lease_token=uuid4(),
    )
    repository = _WorkerRepository(job)
    store = InMemoryObjectStore()
    await store.put_bytes(
        store.material_templates_bucket,
        job.source_object_key or "",
        source,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    worker = MaterialDocumentWorker(
        repository,  # type: ignore[arg-type]
        store,
        MockMaterialTemplateAnalyzer(),
        _SlowPageRenderer(0.05),
        DeterministicDocxMaterialRenderer(),
        lease_seconds=30,
        lease_heartbeat_seconds=0.01,
    )
    assert await worker.run_once() is True
    assert repository.renewals >= 1
    assert repository.failed is None
    assert repository.completed is not None


@pytest.mark.asyncio
async def test_worker_stops_when_lease_renewal_is_rejected() -> None:
    source = _docx()
    job = LeasedMaterialDocumentJob(
        id=uuid4(),
        owner_account_id=uuid4(),
        application_id=uuid4(),
        requirement_code="id-2",
        template_id=uuid4(),
        template_key="id-loss-statement-v1",
        template_title="丢失情况说明",
        template_mode=MaterialTemplateMode.SOURCE_EDITABLE,
        allowed_fields=("applicant_name",),
        source_object_key="material-templates/v1/source.docx",
        source_sha256=hashlib.sha256(source).hexdigest(),
        form_snapshot={"applicant_name": "演示甲"},
        request_text=None,
        lease_token=uuid4(),
    )
    repository = _WorkerRepository(job)
    repository.renew_success = False
    store = InMemoryObjectStore()
    await store.put_bytes(
        store.material_templates_bucket,
        job.source_object_key or "",
        source,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    worker = MaterialDocumentWorker(
        repository,  # type: ignore[arg-type]
        store,
        MockMaterialTemplateAnalyzer(),
        _SlowPageRenderer(0.1),
        DeterministicDocxMaterialRenderer(),
        lease_seconds=30,
        lease_heartbeat_seconds=0.01,
    )
    assert await worker.run_once() is True
    assert repository.renewals == 1
    assert repository.completed is None
    assert repository.failed == "LEASE_LOST"


class _QuotaSession:
    def __init__(self, counts: list[int]) -> None:
        self.counts = iter(counts)
        self.locked = False
        self.scalar_calls = 0
        self.scalar_statements: list[object] = []

    async def execute(self, _statement: object, _parameters: object) -> None:
        self.locked = True

    async def scalar(self, statement: object) -> int:
        assert self.locked
        self.scalar_calls += 1
        self.scalar_statements.append(statement)
        return next(self.counts)


@pytest.mark.asyncio
async def test_worker_lease_scopes_stale_cleanup_and_queue_claim_to_release_lane() -> None:
    statements: list[object] = []

    class Scalars:
        @staticmethod
        def all() -> list[object]:
            return []

    class Session:
        async def scalars(self, statement: object) -> Scalars:
            statements.append(statement)
            return Scalars()

        async def scalar(self, statement: object) -> None:
            statements.append(statement)
            return None

    class Context:
        async def __aenter__(self) -> Session:
            return Session()

        async def __aexit__(self, *_args: object) -> None:
            return None

    class Sessions:
        @staticmethod
        def begin() -> Context:
            return Context()

    repository = BusinessRepository(Sessions())  # type: ignore[arg-type]
    assert await repository.lease_material_document_job(180, "release-a") is None
    assert len(statements) == 2
    for statement in statements:
        assert "release_lane" in str(statement)
        parameters = statement.compile().params  # type: ignore[union-attr]
        assert "release-a" in parameters.values()
        assert "release-b" not in parameters.values()


@pytest.mark.asyncio
async def test_job_creation_rechecks_published_service_under_transaction_lock() -> None:
    owner_id = uuid4()
    application_id = uuid4()
    service_id = uuid4()
    calls: list[tuple[object, object, dict[str, object]]] = []

    class Session:
        async def get(
            self, model: object, identity: object, **kwargs: object
        ) -> object | None:
            calls.append((model, identity, kwargs))
            if model is ApplicationRecord:
                return SimpleNamespace(
                    id=application_id,
                    applicant_id=owner_id,
                    service_id=service_id,
                    status=ApplicationStatus.DRAFT.value,
                )
            if model is GovernmentServiceRecord:
                return SimpleNamespace(
                    id=service_id, status=ServiceStatus.DRAFT.value
                )
            return None

    repository = BusinessRepository(object())  # type: ignore[arg-type]
    session = Session()
    token = repository._uow_session.set(session)  # type: ignore[arg-type]
    try:
        with pytest.raises(ConflictError) as caught:
            await repository.create_material_document_job(
                owner_id,
                application_id,
                "test-1",
                uuid4(),
                None,
                datetime.now(timezone.utc) + timedelta(hours=24),
                "qwen-test-v1",
                "release-test",
                10,
                2,
                20,
                50,
            )
    finally:
        repository._uow_session.reset(token)
    assert caught.value.code == "material_document_service_unavailable"
    assert calls == [
        (ApplicationRecord, application_id, {"with_for_update": True}),
        (GovernmentServiceRecord, service_id, {"with_for_update": True}),
    ]


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("counts", "code"),
    [
        ([10], "material_document_daily_limit_exceeded"),
        ([0, 20], "material_document_global_daily_limit_exceeded"),
        ([0, 0, 2], "material_document_user_active_limit_exceeded"),
        ([0, 0, 0, 50], "material_document_global_queue_limit_exceeded"),
    ],
)
async def test_material_job_admission_limits_are_checked_under_global_lock(
    counts: list[int], code: str
) -> None:
    session = _QuotaSession(counts)
    with pytest.raises(TooManyRequests) as caught:
        await _enforce_material_document_job_limits(
            session,  # type: ignore[arg-type]
            uuid4(),
            datetime.now(timezone.utc),
            user_daily_limit=10,
            user_active_limit=2,
            global_daily_limit=20,
            global_queue_limit=50,
        )
    assert session.locked is True
    assert caught.value.code == code


@pytest.mark.asyncio
async def test_material_job_admission_allows_counts_below_every_limit() -> None:
    session = _QuotaSession([9, 19, 1, 49])
    await _enforce_material_document_job_limits(
        session,  # type: ignore[arg-type]
        uuid4(),
        datetime.now(timezone.utc),
        user_daily_limit=10,
        user_active_limit=2,
        global_daily_limit=20,
        global_queue_limit=50,
    )
    assert session.locked is True
    assert session.scalar_calls == 4


@pytest.mark.asyncio
async def test_global_daily_cap_counts_all_accounts_from_utc_day_start() -> None:
    session = _QuotaSession([0, 0, 0, 0])
    now = datetime(2026, 8, 26, 15, 30, tzinfo=timezone.utc)
    await _enforce_material_document_job_limits(
        session,  # type: ignore[arg-type]
        uuid4(),
        now,
        user_daily_limit=10,
        user_active_limit=2,
        global_daily_limit=20,
        global_queue_limit=50,
    )
    global_daily_statement = session.scalar_statements[1]
    sql = str(global_daily_statement)
    assert "owner_account_id" not in sql
    assert datetime(2026, 8, 26, tzinfo=timezone.utc) in (
        global_daily_statement.compile().params.values()  # type: ignore[union-attr]
    )


@pytest.mark.asyncio
async def test_coordinator_snapshots_model_semantics_and_configured_limits() -> None:
    captured: tuple[object, ...] | None = None

    class Repository:
        async def create_material_document_job(self, *args: object) -> dict[str, object]:
            nonlocal captured
            captured = args
            return {"status": "QUEUED"}

    class Idempotency:
        async def execute(
            self,
            _actor: object,
            _scope: object,
            _key: object,
            _request: object,
            operation: object,
        ) -> dict[str, object]:
            return await operation()  # type: ignore[operator]

    coordinator = MaterialDocumentCoordinator(
        Repository(),  # type: ignore[arg-type]
        InMemoryObjectStore(),
        Idempotency(),
        enabled=True,
        retention_hours=24,
        model_name="qwen-test-v1",
        release_lane="release-test",
        user_daily_limit=7,
        user_active_limit=2,
        global_daily_limit=19,
        global_queue_limit=31,
    )
    await coordinator.create(
        _citizen(), uuid4(), "id-2", uuid4(), "  演示需求  ", "idem-1"
    )
    assert captured is not None
    assert captured[-6:] == ("qwen-test-v1", "release-test", 7, 2, 19, 31)


@pytest.mark.asyncio
async def test_status_never_exposes_private_object_key() -> None:
    principal = _citizen()
    generation_id = uuid4()

    class Repository:
        async def get_material_document_job_authorized(
            self, _generation_id: object, _owner: object
        ) -> dict[str, object]:
            return {
                "generation_id": generation_id,
                "status": "READY",
                "object_key": "private/secret.docx",
                "expires_at": datetime.now(timezone.utc) + timedelta(hours=1),
            }

    coordinator = MaterialDocumentCoordinator(
        Repository(),  # type: ignore[arg-type]
        InMemoryObjectStore(),
        object(),
        enabled=True,
        retention_hours=24,
    )
    status = await coordinator.status(principal, generation_id)
    assert "object_key" not in status


@pytest.mark.asyncio
async def test_generated_bucket_alone_gets_one_day_lifecycle() -> None:
    class Client:
        lifecycle: tuple[str, object] | None = None

        @staticmethod
        def bucket_exists(_bucket: str) -> bool:
            return True

        def set_bucket_lifecycle(self, bucket: str, config: object) -> None:
            self.lifecycle = (bucket, config)

    store = object.__new__(MinioObjectStore)
    store._client = Client()  # type: ignore[attr-defined]
    store.materials_bucket = "materials"
    store.knowledge_bucket = "knowledge"
    store.material_templates_bucket = "templates"
    store.generated_documents_bucket = "generated"
    await store.ensure_buckets()
    bucket, config = store._client.lifecycle  # type: ignore[misc,union-attr]
    assert bucket == "generated"
    assert config.rules[0].expiration.days == 1  # type: ignore[attr-defined]
    assert config.rules[0].rule_filter.prefix == ""  # type: ignore[attr-defined]
